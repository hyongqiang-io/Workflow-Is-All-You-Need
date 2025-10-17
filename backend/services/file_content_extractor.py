"""
文件内容提取器
File Content Extractor - Linus式简洁设计

专为工作流系统中的AI任务附件理解而设计
支持多种文件格式的内容提取，转换为AI可理解的文本格式
"""

import os
import asyncio
import mimetypes
from pathlib import Path
from typing import Dict, Any, Optional
from loguru import logger


class FileContentExtractor:
    """文件内容提取器 - 统一接口设计"""

    def __init__(self):
        """初始化提取器"""
        # 支持的文件格式映射
        self.extractors = {
            # PDF文档
            'pdf': self._extract_pdf,
            'application/pdf': self._extract_pdf,

            # 文本文件
            'txt': self._extract_text,
            'text/plain': self._extract_text,
            'md': self._extract_text,
            'text/markdown': self._extract_text,

            # Office文档
            'docx': self._extract_docx,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx,
            'doc': self._extract_doc_fallback,
            'application/msword': self._extract_doc_fallback,

            # Excel文件
            'xlsx': self._extract_excel,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_excel,
            'xls': self._extract_excel_fallback,
            'application/vnd.ms-excel': self._extract_excel_fallback,

            # 图片文件
            'jpg': self._extract_image,
            'jpeg': self._extract_image,
            'png': self._extract_image,
            'gif': self._extract_image,
            'bmp': self._extract_image,
            'tiff': self._extract_image,
            'image/jpeg': self._extract_image,
            'image/png': self._extract_image,
            'image/gif': self._extract_image,
            'image/bmp': self._extract_image,
            'image/tiff': self._extract_image,
        }

        # 初始化状态
        self._dependencies_checked = False
        self._available_extractors = set()

    async def extract_content(self, file_path: str, content_type: Optional[str] = None) -> Dict[str, Any]:
        """
        统一的文件内容提取接口

        Args:
            file_path: 文件路径
            content_type: MIME类型（可选，会自动检测）

        Returns:
            提取结果字典，包含success, content, metadata等字段
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                return self._error_result(f"文件不存在: {file_path}")

            # 确定文件类型
            file_type = self._determine_file_type(file_path, content_type)

            # 检查是否支持该文件类型
            if file_type not in self.extractors:
                return self._unsupported_result(file_path, file_type)

            # 执行内容提取
            extractor = self.extractors[file_type]
            result = await extractor(file_path)

            # 统一返回格式
            return {
                'success': True,
                'content': result.get('text', ''),
                'metadata': {
                    'file_path': file_path,
                    'file_type': file_type,
                    'file_size': os.path.getsize(file_path),
                    'extractor': result.get('extractor', 'unknown'),
                    **result.get('metadata', {})
                }
            }

        except Exception as e:
            logger.error(f"文件内容提取失败: {file_path}, 错误: {e}")
            return self._error_result(f"提取失败: {str(e)}")

    def _determine_file_type(self, file_path: str, content_type: Optional[str] = None) -> str:
        """确定文件类型"""
        # 优先使用提供的content_type
        if content_type and content_type in self.extractors:
            return content_type

        # 基于文件扩展名
        file_ext = Path(file_path).suffix.lower().lstrip('.')
        if file_ext in self.extractors:
            return file_ext

        # 使用mimetypes模块检测
        detected_type, _ = mimetypes.guess_type(file_path)
        if detected_type and detected_type in self.extractors:
            return detected_type

        # 默认返回扩展名
        return file_ext or 'unknown'

    def _error_result(self, error_msg: str) -> Dict[str, Any]:
        """生成错误结果"""
        return {
            'success': False,
            'content': f"[错误: {error_msg}]",
            'error': error_msg,
            'metadata': {}
        }

    def _unsupported_result(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """生成不支持格式的结果"""
        filename = os.path.basename(file_path)
        return {
            'success': False,
            'content': f"[不支持的文件格式: {filename} ({file_type})]",
            'error': f"Unsupported file type: {file_type}",
            'metadata': {'file_type': file_type}
        }

    # ==================== 具体的提取器实现 ====================

    async def _extract_text(self, file_path: str) -> Dict[str, Any]:
        """提取纯文本文件内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()

                    return {
                        'text': content,
                        'extractor': 'text_reader',
                        'metadata': {'encoding': encoding}
                    }
                except UnicodeDecodeError:
                    continue

            # 所有编码都失败，返回错误
            raise ValueError("无法确定文件编码")

        except Exception as e:
            logger.error(f"文本文件提取失败: {file_path}, 错误: {e}")
            raise

    async def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """提取PDF文件内容 - 优先使用pymupdf4llm，降级到PyPDF2"""
        try:
            # 方案1: 优先使用pymupdf4llm (最佳质量)
            try:
                import pymupdf4llm

                # 转换为Markdown格式，保留文档结构
                markdown_text = pymupdf4llm.to_markdown(file_path)

                # 基本的内容验证
                if len(markdown_text.strip()) < 10:
                    raise ValueError("PDF内容提取结果过短，可能提取失败")

                return {
                    'text': markdown_text,
                    'extractor': 'pymupdf4llm',
                    'metadata': {
                        'format': 'markdown',
                        'content_length': len(markdown_text),
                        'quality': 'high'
                    }
                }

            except ImportError:
                logger.info("pymupdf4llm未安装，使用PyPDF2降级方案")
                return await self._extract_pdf_fallback(file_path)
            except Exception as e:
                logger.warning(f"pymupdf4llm提取失败: {e}，使用PyPDF2降级方案")
                return await self._extract_pdf_fallback(file_path)

        except Exception as e:
            logger.error(f"PDF提取失败: {file_path}, 错误: {e}")
            raise

    async def _extract_pdf_fallback(self, file_path: str) -> Dict[str, Any]:
        """PDF降级提取方案 - 使用PyPDF2"""
        try:
            import PyPDF2

            text_content = []

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # 检查PDF是否被加密
                if pdf_reader.is_encrypted:
                    # 尝试用空密码解密
                    try:
                        pdf_reader.decrypt("")
                    except:
                        return {
                            'text': f"[PDF文档已加密，无法提取内容: {os.path.basename(file_path)}]",
                            'extractor': 'pypdf2_encrypted',
                            'metadata': {'encrypted': True}
                        }

                # 提取每一页的文本
                total_pages = len(pdf_reader.pages)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- 第{page_num + 1}页 ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"提取PDF第{page_num + 1}页失败: {e}")
                        text_content.append(f"--- 第{page_num + 1}页 ---\n[页面提取失败]")

                final_text = "\n\n".join(text_content)

                # 基本的内容验证
                if len(final_text.strip()) < 10:
                    return {
                        'text': f"[PDF文档内容提取失败或为空: {os.path.basename(file_path)}]",
                        'extractor': 'pypdf2_empty',
                        'metadata': {'total_pages': total_pages, 'extracted_length': len(final_text)}
                    }

                return {
                    'text': final_text,
                    'extractor': 'pypdf2',
                    'metadata': {
                        'format': 'text',
                        'total_pages': total_pages,
                        'content_length': len(final_text),
                        'quality': 'basic'
                    }
                }

        except ImportError:
            return {
                'text': f"[PDF处理库未安装，无法提取内容: {os.path.basename(file_path)}]",
                'extractor': 'no_pdf_library',
                'metadata': {'missing_dependency': 'PyPDF2'}
            }
        except Exception as e:
            logger.error(f"PyPDF2提取失败: {file_path}, 错误: {e}")
            return {
                'text': f"[PDF处理失败: {os.path.basename(file_path)} - {str(e)}]",
                'extractor': 'pypdf2_error',
                'metadata': {'error': str(e)}
            }

    async def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """提取Word文档内容"""
        try:
            import docx

            doc = docx.Document(file_path)
            content_parts = []

            # 提取段落内容
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            # 提取表格内容
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_content.append(cell_text if cell_text else "-")
                    table_content.append(" | ".join(row_content))

                if table_content:
                    content_parts.append("\n表格内容:")
                    content_parts.extend(table_content)

            final_text = "\n\n".join(content_parts)

            if len(final_text.strip()) < 5:
                return {
                    'text': f"[Word文档为空或无可提取内容: {os.path.basename(file_path)}]",
                    'extractor': 'docx_empty',
                    'metadata': {'paragraphs': len(doc.paragraphs), 'tables': len(doc.tables)}
                }

            return {
                'text': final_text,
                'extractor': 'python-docx',
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'content_length': len(final_text)
                }
            }

        except ImportError:
            return {
                'text': f"[Word文档处理库未安装，无法提取内容: {os.path.basename(file_path)}]",
                'extractor': 'no_docx_library',
                'metadata': {'missing_dependency': 'python-docx'}
            }
        except Exception as e:
            logger.error(f"Word文档提取失败: {file_path}, 错误: {e}")
            return {
                'text': f"[Word文档处理失败: {os.path.basename(file_path)} - {str(e)}]",
                'extractor': 'docx_error',
                'metadata': {'error': str(e)}
            }

    async def _extract_doc_fallback(self, file_path: str) -> Dict[str, Any]:
        """提取旧版Word文档内容（降级方案）"""
        return {
            'text': f"[旧版Word文档: {os.path.basename(file_path)}，建议转换为docx格式]",
            'extractor': 'doc_fallback',
            'metadata': {'format': 'doc_legacy'}
        }

    async def _extract_excel(self, file_path: str) -> Dict[str, Any]:
        """提取Excel文件内容"""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # 跳过空工作表
                if sheet.max_row == 1 and sheet.max_column == 1:
                    continue

                content_parts.append(f"=== 工作表: {sheet_name} ===")

                # 提取工作表数据
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空行
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):
                        content_parts.append(" | ".join(row_data))

            final_text = "\n".join(content_parts)

            if len(final_text.strip()) < 10:
                return {
                    'text': f"[Excel文档为空或无可提取内容: {os.path.basename(file_path)}]",
                    'extractor': 'excel_empty',
                    'metadata': {'worksheets': len(workbook.sheetnames)}
                }

            return {
                'text': final_text,
                'extractor': 'openpyxl',
                'metadata': {
                    'worksheets': len(workbook.sheetnames),
                    'content_length': len(final_text)
                }
            }

        except ImportError:
            return {
                'text': f"[Excel处理库未安装，无法提取内容: {os.path.basename(file_path)}]",
                'extractor': 'no_excel_library',
                'metadata': {'missing_dependency': 'openpyxl'}
            }
        except Exception as e:
            logger.error(f"Excel文件提取失败: {file_path}, 错误: {e}")
            return {
                'text': f"[Excel文件处理失败: {os.path.basename(file_path)} - {str(e)}]",
                'extractor': 'excel_error',
                'metadata': {'error': str(e)}
            }

    async def _extract_excel_fallback(self, file_path: str) -> Dict[str, Any]:
        """提取旧版Excel文件内容（降级方案）"""
        return {
            'text': f"[旧版Excel文档: {os.path.basename(file_path)}，建议转换为xlsx格式]",
            'extractor': 'excel_fallback',
            'metadata': {'format': 'excel_legacy'}
        }

    async def _extract_image(self, file_path: str) -> Dict[str, Any]:
        """提取图片文件内容 - 可选OCR功能"""
        try:
            # 方案1: 尝试使用OCR提取图片中的文字
            try:
                from PIL import Image
                import pytesseract

                # 打开图片
                image = Image.open(file_path)

                # 使用Tesseract进行OCR
                # 支持中英文混合识别
                ocr_text = pytesseract.image_to_string(
                    image,
                    lang='chi_sim+eng',  # 中文简体 + 英文
                    config='--psm 6'     # 页面分割模式
                )

                # 清理OCR结果
                cleaned_text = ocr_text.strip()

                if len(cleaned_text) < 3:
                    # OCR没有识别到有效文字
                    return self._image_placeholder_result(file_path, "OCR未识别到文字内容")

                return {
                    'text': f"[图片OCR内容]\n{cleaned_text}",
                    'extractor': 'tesseract_ocr',
                    'metadata': {
                        'image_size': image.size,
                        'image_mode': image.mode,
                        'ocr_length': len(cleaned_text),
                        'languages': 'chi_sim+eng'
                    }
                }

            except ImportError as e:
                missing_lib = 'pytesseract' if 'pytesseract' in str(e) else 'Pillow'
                return self._image_placeholder_result(file_path, f"OCR库未安装: {missing_lib}")

            except Exception as e:
                logger.warning(f"图片OCR失败: {file_path}, 错误: {e}")
                return self._image_placeholder_result(file_path, f"OCR处理失败: {str(e)}")

        except Exception as e:
            logger.error(f"图片处理失败: {file_path}, 错误: {e}")
            return self._image_placeholder_result(file_path, f"图片处理错误: {str(e)}")

    def _image_placeholder_result(self, file_path: str, reason: str) -> Dict[str, Any]:
        """生成图片占位符结果"""
        filename = os.path.basename(file_path)
        return {
            'text': f"[图片文件: {filename} - {reason}]",
            'extractor': 'image_placeholder',
            'metadata': {
                'placeholder_reason': reason,
                'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
            }
        }

    async def extract_preview_content(self, file_path: str, content_type: str, max_size: int = 1024 * 1024) -> Dict[str, Any]:
        """
        为预览功能提取文件内容 - Linus式简洁设计

        Args:
            file_path: 文件路径
            content_type: MIME类型
            max_size: 最大预览大小限制

        Returns:
            预览数据字典，包含preview_type, content等字段
        """
        try:
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size > max_size:
                return {
                    'preview_type': 'error',
                    'content': f'文件过大，无法预览 (大小: {file_size / 1024 / 1024:.1f}MB，限制: {max_size / 1024 / 1024:.1f}MB)',
                    'error': 'file_too_large'
                }

            # 根据内容类型决定预览方式
            if self._is_text_type(content_type):
                return await self._preview_as_text(file_path, content_type)
            elif self._is_image_type(content_type):
                return await self._preview_as_image(file_path, content_type)
            elif self._is_pdf_type(content_type):
                return await self._preview_as_pdf(file_path)
            elif self._is_office_type(content_type):
                return await self._preview_as_office(file_path, content_type)
            else:
                return await self._preview_as_metadata(file_path, content_type)

        except Exception as e:
            logger.error(f"文件预览提取失败: {file_path}, 错误: {e}")
            return {
                'preview_type': 'error',
                'content': f'预览失败: {str(e)}',
                'error': str(e)
            }

    def _is_text_type(self, content_type: str) -> bool:
        """判断是否为文本类型文件"""
        text_types = [
            'text/plain', 'text/markdown', 'text/csv', 'text/html',
            'application/json', 'application/xml'
        ]
        return content_type in text_types or content_type.startswith('text/')

    def _is_image_type(self, content_type: str) -> bool:
        """判断是否为图片类型文件"""
        return content_type.startswith('image/')

    def _is_pdf_type(self, content_type: str) -> bool:
        """判断是否为PDF文件"""
        return content_type == 'application/pdf'

    def _is_office_type(self, content_type: str) -> bool:
        """判断是否为Office文档"""
        office_types = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            'application/msword',
            'application/vnd.ms-excel',
            'application/vnd.ms-powerpoint'
        ]
        return content_type in office_types

    async def _preview_as_text(self, file_path: str, content_type: str) -> Dict[str, Any]:
        """以文本形式预览"""
        try:
            # 限制预览长度
            max_chars = 50000  # 50KB文本

            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        content = f.read(max_chars)

                    return {
                        'preview_type': 'text',
                        'content': content,
                        'encoding': encoding,
                        'truncated': len(content) >= max_chars
                    }
                except UnicodeDecodeError:
                    continue

            return {
                'preview_type': 'error',
                'content': '无法识别文件编码',
                'error': 'encoding_error'
            }

        except Exception as e:
            return {
                'preview_type': 'error',
                'content': f'文本预览失败: {str(e)}',
                'error': str(e)
            }

    async def _preview_as_image(self, file_path: str, content_type: str) -> Dict[str, Any]:
        """以图片形式预览 - 返回base64数据"""
        try:
            import base64

            # 限制图片预览大小
            max_image_size = 5 * 1024 * 1024  # 5MB
            file_size = os.path.getsize(file_path)

            if file_size > max_image_size:
                return {
                    'preview_type': 'error',
                    'content': f'图片过大，无法预览 ({file_size / 1024 / 1024:.1f}MB)',
                    'error': 'image_too_large'
                }

            with open(file_path, 'rb') as f:
                image_data = f.read()

            base64_data = base64.b64encode(image_data).decode('utf-8')

            return {
                'preview_type': 'image',
                'content': base64_data,
                'content_type': content_type,
                'file_size': file_size
            }

        except Exception as e:
            return {
                'preview_type': 'error',
                'content': f'图片预览失败: {str(e)}',
                'error': str(e)
            }

    async def _preview_as_pdf(self, file_path: str) -> Dict[str, Any]:
        """PDF预览 - 提取前几页文本"""
        try:
            result = await self._extract_pdf(file_path)

            # 限制预览长度
            content = result.get('text', '')
            max_chars = 10000  # 10KB文本预览

            if len(content) > max_chars:
                content = content[:max_chars] + '\n\n[内容已截断，完整内容请下载查看]'

            return {
                'preview_type': 'pdf_text',
                'content': content,
                'extractor': result.get('extractor', 'pdf'),
                'metadata': result.get('metadata', {}),
                'truncated': len(result.get('text', '')) > max_chars
            }

        except Exception as e:
            return {
                'preview_type': 'error',
                'content': f'PDF预览失败: {str(e)}',
                'error': str(e)
            }

    async def _preview_as_office(self, file_path: str, content_type: str) -> Dict[str, Any]:
        """Office文档预览"""
        try:
            # 根据文档类型选择提取器
            if 'wordprocessingml' in content_type:
                result = await self._extract_docx(file_path)
            elif 'spreadsheetml' in content_type:
                result = await self._extract_excel(file_path)
            else:
                result = await self._extract_docx(file_path)  # 默认用docx提取器

            # 限制预览长度
            content = result.get('text', '')
            max_chars = 10000  # 10KB文本预览

            if len(content) > max_chars:
                content = content[:max_chars] + '\n\n[内容已截断，完整内容请下载查看]'

            return {
                'preview_type': 'office_text',
                'content': content,
                'extractor': result.get('extractor', 'office'),
                'metadata': result.get('metadata', {}),
                'truncated': len(result.get('text', '')) > max_chars
            }

        except Exception as e:
            return {
                'preview_type': 'error',
                'content': f'Office文档预览失败: {str(e)}',
                'error': str(e)
            }

    async def _preview_as_metadata(self, file_path: str, content_type: str) -> Dict[str, Any]:
        """其他文件类型只返回元数据"""
        try:
            file_stat = os.stat(file_path)

            return {
                'preview_type': 'metadata',
                'content': {
                    'filename': os.path.basename(file_path),
                    'content_type': content_type,
                    'file_size': file_stat.st_size,
                    'created_time': file_stat.st_ctime,
                    'modified_time': file_stat.st_mtime,
                    'message': '此文件类型不支持内容预览，请下载查看'
                }
            }

        except Exception as e:
            return {
                'preview_type': 'error',
                'content': f'获取文件信息失败: {str(e)}',
                'error': str(e)
            }


    # ==================== 多模态AI支持方法 ====================

    async def extract_content_for_multimodal(self, file_path: str, content_type: Optional[str] = None) -> Dict[str, Any]:
        """
        为多模态AI提取内容，图片返回base64编码

        Args:
            file_path: 文件路径
            content_type: MIME类型

        Returns:
            包含content、content_type、is_image等信息的字典
        """
        try:
            # 确定文件类型
            if not content_type:
                content_type, _ = mimetypes.guess_type(file_path)

            file_ext = Path(file_path).suffix.lower().lstrip('.')

            # 判断是否为图片文件
            image_extensions = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp']
            image_mime_types = ['image/jpeg', 'image/png', 'image/gif', 'image/bmp', 'image/tiff', 'image/webp']

            is_image = (file_ext in image_extensions) or (content_type in image_mime_types)

            if is_image:
                # 图片文件：返回base64编码
                return await self._extract_image_as_base64(file_path, content_type)
            else:
                # 非图片文件：使用常规文本提取
                result = await self.extract_content(file_path, content_type)
                return {
                    'success': result['success'],
                    'content': result.get('content', ''),
                    'content_type': 'text',
                    'is_image': False,
                    'extractor': result.get('extractor', 'text'),
                    'error': result.get('error')
                }

        except Exception as e:
            logger.error(f"多模态内容提取失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'content': '',
                'content_type': 'text',
                'is_image': False,
                'error': f"多模态内容提取失败: {str(e)}"
            }

    async def _extract_image_as_base64(self, file_path: str, content_type: Optional[str] = None) -> Dict[str, Any]:
        """
        将图片文件转换为base64编码

        Args:
            file_path: 图片文件路径
            content_type: MIME类型

        Returns:
            包含base64编码的结果字典
        """
        try:
            import base64

            # 读取文件二进制数据
            with open(file_path, 'rb') as f:
                image_data = f.read()

            # 转换为base64
            base64_data = base64.b64encode(image_data).decode('utf-8')

            # 确定MIME类型
            if not content_type:
                content_type, _ = mimetypes.guess_type(file_path)
                if not content_type:
                    # 根据文件扩展名推断
                    ext = Path(file_path).suffix.lower()
                    mime_map = {
                        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                        '.png': 'image/png', '.gif': 'image/gif',
                        '.bmp': 'image/bmp', '.tiff': 'image/tiff',
                        '.webp': 'image/webp'
                    }
                    content_type = mime_map.get(ext, 'image/jpeg')

            # 获取图片基本信息
            file_size = len(image_data)
            file_name = os.path.basename(file_path)

            logger.info(f"✅ [MULTIMODAL] 图片转base64成功: {file_name}, 大小: {file_size} bytes")

            return {
                'success': True,
                'content': base64_data,
                'content_type': content_type,
                'is_image': True,
                'extractor': 'base64_encoder',
                'metadata': {
                    'file_name': file_name,
                    'file_size': file_size,
                    'mime_type': content_type,
                    'encoding': 'base64'
                }
            }

        except Exception as e:
            logger.error(f"❌ [MULTIMODAL] 图片base64转换失败: {file_path}, 错误: {e}")
            return {
                'success': False,
                'content': f"[图片处理失败: {os.path.basename(file_path)}]",
                'content_type': 'text',
                'is_image': False,
                'error': f"图片base64转换失败: {str(e)}"
            }

    # ==================== 工具方法 ====================

    def get_supported_formats(self) -> Dict[str, list]:
        """获取支持的文件格式列表"""
        formats = {
            'documents': ['pdf', 'docx', 'doc', 'txt', 'md'],
            'spreadsheets': ['xlsx', 'xls'],
            'images': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff'],
            'mime_types': [key for key in self.extractors.keys() if '/' in key]
        }
        return formats

    async def check_dependencies(self) -> Dict[str, bool]:
        """检查依赖库的可用性"""
        if self._dependencies_checked:
            return {'cached': True}

        dependencies = {}

        # 检查pymupdf4llm
        try:
            import pymupdf4llm
            dependencies['pymupdf4llm'] = True
            self._available_extractors.add('pdf_advanced')
        except ImportError:
            dependencies['pymupdf4llm'] = False

        # 检查PyPDF2
        try:
            import PyPDF2
            dependencies['PyPDF2'] = True
            self._available_extractors.add('pdf_basic')
        except ImportError:
            dependencies['PyPDF2'] = False

        # 检查python-docx
        try:
            import docx
            dependencies['python-docx'] = True
            self._available_extractors.add('docx')
        except ImportError:
            dependencies['python-docx'] = False

        # 检查openpyxl
        try:
            import openpyxl
            dependencies['openpyxl'] = True
            self._available_extractors.add('excel')
        except ImportError:
            dependencies['openpyxl'] = False

        # 检查PIL和pytesseract
        try:
            from PIL import Image
            dependencies['pillow'] = True
            try:
                import pytesseract
                dependencies['pytesseract'] = True
                self._available_extractors.add('ocr')
            except ImportError:
                dependencies['pytesseract'] = False
        except ImportError:
            dependencies['pillow'] = False
            dependencies['pytesseract'] = False

        self._dependencies_checked = True
        return dependencies

    async def extract_task_attachments(self, task_id) -> str:
        """
        提取任务的所有附件内容并整合
        Args:
            task_id: 任务实例ID
        Returns:
            整合后的附件内容文本
        """
        try:
            import uuid
            from .file_association_service import FileAssociationService

            logger.info(f"📎 [ATTACHMENT-EXTRACT] 开始提取任务附件: {task_id}")

            # 获取任务相关的文件
            file_service = FileAssociationService()

            # 1. 首先查询直接关联的任务附件 (task_instance_file)
            task_files = await file_service.get_task_instance_files(uuid.UUID(str(task_id)))
            logger.info(f"📎 [ATTACHMENT-EXTRACT] 直接任务附件: {len(task_files) if task_files else 0} 个")

            # 2. 如果没有直接任务附件，查询节点级别的附件 (node_instance_file)
            node_files = []
            if not task_files:
                try:
                    # 获取任务的node_instance_id
                    from ..repositories.instance.task_instance_repository import TaskInstanceRepository
                    task_repo = TaskInstanceRepository()
                    task_info = await task_repo.get_task_by_id(uuid.UUID(str(task_id)))

                    if task_info and task_info.get('node_instance_id'):
                        node_instance_id = task_info['node_instance_id']
                        logger.info(f"📎 [ATTACHMENT-EXTRACT] 查询节点附件，node_instance_id: {node_instance_id}")

                        # 查询节点级别的文件
                        node_files = await file_service.get_node_instance_files(uuid.UUID(str(node_instance_id)))
                        logger.info(f"📎 [ATTACHMENT-EXTRACT] 节点级别附件: {len(node_files) if node_files else 0} 个")

                except Exception as e:
                    logger.warning(f"⚠️ [ATTACHMENT-EXTRACT] 查询节点附件失败: {e}")

            # 3. 合并所有文件
            all_files = (task_files or []) + (node_files or [])

            if not all_files:
                logger.info(f"📎 [ATTACHMENT-EXTRACT] 任务 {task_id} 无任何附件")
                return ""

            logger.info(f"📎 [ATTACHMENT-EXTRACT] 总共找到 {len(all_files)} 个附件")

            all_content = []
            for file_info in all_files:
                try:
                    file_path = file_info.get('file_path', '')
                    file_name = file_info.get('file_name', '') or file_info.get('original_filename', 'unknown')

                    logger.info(f"📄 [ATTACHMENT-EXTRACT] 处理文件: {file_name}")

                    if os.path.exists(file_path):
                        # 提取文件内容
                        result = await self.extract_content(file_path)

                        if result['success'] and result['content']:
                            content_header = f"\n{'='*50}\n文件: {file_name}\n路径: {file_path}\n{'='*50}\n"
                            all_content.append(content_header + result['content'])
                            logger.info(f"✅ [ATTACHMENT-EXTRACT] 文件 {file_name} 提取成功")
                        else:
                            logger.warning(f"⚠️ [ATTACHMENT-EXTRACT] 文件 {file_name} 提取失败: {result.get('error', 'unknown')}")
                    else:
                        logger.warning(f"⚠️ [ATTACHMENT-EXTRACT] 文件不存在: {file_path}")

                except Exception as e:
                    logger.error(f"❌ [ATTACHMENT-EXTRACT] 处理文件失败: {e}")
                    continue

            combined_content = "\n\n".join(all_content)

            if combined_content:
                logger.info(f"✅ [ATTACHMENT-EXTRACT] 附件内容提取完成，总长度: {len(combined_content)} 字符")
            else:
                logger.info(f"📎 [ATTACHMENT-EXTRACT] 无有效附件内容")

            return combined_content

        except Exception as e:
            logger.error(f"❌ [ATTACHMENT-EXTRACT] 提取任务附件失败: {e}")
            import traceback
            logger.error(f"   错误堆栈: {traceback.format_exc()}")
            return f"提取附件内容时出错: {str(e)}"


# 全局实例
file_content_extractor = FileContentExtractor()

def get_file_content_extractor() -> FileContentExtractor:
    """获取文件内容提取器实例"""
    return file_content_extractor